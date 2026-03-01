"""
AlumniTrust AI — FastAPI Backend
Higher-Ed Security Questionnaire Automation with RAG
"""

import os
import io
import tempfile
from datetime import datetime, timedelta, timezone
from typing import List, Optional

import jwt
import pandas as pd
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from fastapi import (
    FastAPI, Depends, HTTPException, UploadFile,
    File, Form, BackgroundTasks, Header
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from passlib.context import CryptContext
from pydantic import BaseModel
from sqlalchemy.orm import Session

from models import User, Document, Task, Answer, get_db, SessionLocal, engine, Base
from engine import ingest_document, batch_process_questionnaire, delete_user_documents

# ─────────────────────────────────────────────
# App bootstrap
# ─────────────────────────────────────────────
from contextlib import asynccontextmanager

@asynccontextmanager
async def lifespan(app: FastAPI):
    try:
        from seed_demo_data import seed_demo_user, seed_documents
        print("[System] Running auto-seeder for demo data (Render compatibility)...")
        user_id = seed_demo_user()
        seed_documents(user_id)
        print("[System] Auto-seeder finished.")
    except Exception as e:
        print(f"[System] Warning: Auto-seeder skipped/failed: {e}")
    yield

app = FastAPI(
    title="AlumniTrust AI",
    description="HECVAT Automation with RAG — powered by Almabase",
    version="1.0.0",
    lifespan=lifespan
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

Base.metadata.create_all(bind=engine)

# ─────────────────────────────────────────────
# Security
# ─────────────────────────────────────────────
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
SECRET_KEY = os.getenv("SECRET_KEY", "alumni-trust-secret-key-changeme-in-prod")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 1440  # 24 h


def hash_password(password: str) -> str:
    return pwd_context.hash(password)


def verify_password(plain: str, hashed: str) -> bool:
    return pwd_context.verify(plain, hashed)


def create_access_token(data: dict, expires_delta: Optional[timedelta] = None) -> str:
    payload = data.copy()
    expire = datetime.now(timezone.utc) + (
        expires_delta if expires_delta else timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    )
    payload["exp"] = expire
    return jwt.encode(payload, SECRET_KEY, algorithm=ALGORITHM)


def _decode_token(token: str) -> int:
    """Decode a JWT and return the user_id (sub). Raises 401 on failure."""
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id = payload.get("sub")
        if user_id is None:
            raise HTTPException(status_code=401, detail="Invalid token")
        return int(user_id)
    except jwt.PyJWTError:
        raise HTTPException(status_code=401, detail="Could not validate credentials")


def get_current_user(
    authorization: Optional[str] = Header(None),
    db: Session = Depends(get_db)
) -> User:
    """
    Extract the current user from the Authorization: Bearer <token> header.
    All protected endpoints use this dependency.
    """
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Not authenticated")
    token = authorization.split(" ", 1)[1]
    user_id = _decode_token(token)
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    return user


# ─────────────────────────────────────────────
# Pydantic schemas
# ─────────────────────────────────────────────

class UserRegister(BaseModel):
    username: str
    email: str
    password: str


class UserLogin(BaseModel):
    email: str
    password: str


class TokenResponse(BaseModel):
    access_token: str
    token_type: str
    user_id: int
    username: str


class DocumentResponse(BaseModel):
    id: int
    filename: str
    created_at: datetime

    class Config:
        from_attributes = True


class AnswerResponse(BaseModel):
    id: int
    question: str
    ai_answer: str
    confidence_score: float
    source_document: str
    source_snippet: str
    has_been_edited: bool
    manual_answer: Optional[str]

    class Config:
        from_attributes = True


class AnswerUpdate(BaseModel):
    manual_answer: str


class TaskSummaryResponse(BaseModel):
    id: int
    task_name: str
    status: str
    created_at: datetime
    answer_count: int

    class Config:
        from_attributes = True


class TaskResponse(BaseModel):
    id: int
    task_name: str
    status: str
    created_at: datetime
    answers: List[AnswerResponse]

    class Config:
        from_attributes = True


# ─────────────────────────────────────────────
# Auth endpoints
# ─────────────────────────────────────────────

@app.post("/api/auth/register", response_model=TokenResponse, tags=["Auth"])
def register(user_data: UserRegister, db: Session = Depends(get_db)):
    """Register a new user and return a JWT token."""
    existing = db.query(User).filter(
        (User.email == user_data.email) | (User.username == user_data.username)
    ).first()
    if existing:
        raise HTTPException(status_code=400, detail="User already exists")

    user = User(
        username=user_data.username,
        email=user_data.email,
        hashed_password=hash_password(user_data.password)
    )
    db.add(user)
    db.commit()
    db.refresh(user)

    token = create_access_token({"sub": user.id})
    return TokenResponse(
        access_token=token, token_type="bearer",
        user_id=user.id, username=user.username
    )


@app.post("/api/auth/login", response_model=TokenResponse, tags=["Auth"])
def login(user_data: UserLogin, db: Session = Depends(get_db)):
    """Login and return a JWT token."""
    user = db.query(User).filter(User.email == user_data.email).first()
    if not user or not verify_password(user_data.password, user.hashed_password):
        raise HTTPException(status_code=401, detail="Invalid email or password")

    token = create_access_token({"sub": user.id})
    return TokenResponse(
        access_token=token, token_type="bearer",
        user_id=user.id, username=user.username
    )


# ─────────────────────────────────────────────
# Document endpoints  (Trust Center Knowledge Base)
# ─────────────────────────────────────────────

@app.post("/api/documents/upload", tags=["Documents"])
async def upload_document(
    file: UploadFile = File(...),
    token: str = Form(...),          # multipart sends token as a form field
    db: Session = Depends(get_db)
):
    """Upload and ingest a reference document into the Trust Center Knowledge Base."""
    user_id = _decode_token(token)
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=401, detail="User not found")

    content_bytes = await file.read()
    file_text = content_bytes.decode("utf-8", errors="ignore")

    doc = Document(
        user_id=user.id,
        filename=file.filename,
        file_path=f"/uploads/{file.filename}",
        content=file_text
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    try:
        ingest_document(file_text, file.filename, user.id, doc.id)
    except Exception as e:
        db.delete(doc)
        db.commit()
        raise HTTPException(status_code=500, detail=f"Error ingesting document: {e}")

    return {
        "message": "Document uploaded successfully to Trust Center Knowledge Base",
        "document_id": doc.id,
        "filename": doc.filename
    }


@app.get("/api/documents", tags=["Documents"])
async def list_documents(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List all reference documents for the authenticated user."""
    docs = db.query(Document).filter(Document.user_id == current_user.id).all()
    return [DocumentResponse.from_orm(d) for d in docs]


@app.delete("/api/documents/{doc_id}", tags=["Documents"])
async def delete_document(
    doc_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Remove a document from the knowledge base and from ChromaDB."""
    doc = db.query(Document).filter(
        Document.id == doc_id, Document.user_id == current_user.id
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    delete_user_documents(current_user.id, doc_id)
    db.delete(doc)
    db.commit()
    return {"message": f"Document '{doc.filename}' deleted"}


# ─────────────────────────────────────────────
# Questionnaire endpoints
# ─────────────────────────────────────────────

@app.post("/api/questionnaire/upload", tags=["Questionnaire"])
async def upload_questionnaire(
    file: UploadFile = File(...),
    task_name: str = Form(...),
    token: str = Form(...),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    db: Session = Depends(get_db)
):
    """
    Upload a CSV or XLSX questionnaire.
    Processing runs in the background; poll GET /api/tasks/{task_id} for status.
    """
    user_id = _decode_token(token)
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=401, detail="User not found")

    file_content = await file.read()

    task = Task(
        user_id=user.id,
        task_name=task_name,
        file_path=f"/uploads/questionnaire_{datetime.utcnow().timestamp()}_{file.filename}",
        status="processing"
    )
    db.add(task)
    db.commit()
    db.refresh(task)

    # Pass task_id + user_id only; background task creates its own DB session
    background_tasks.add_task(
        _bg_process_questionnaire,
        user.id, task.id, file_content, file.filename
    )

    return {"task_id": task.id, "message": "Processing started", "status": "processing"}


def _bg_process_questionnaire(
    user_id: int, task_id: int, file_content: bytes, filename: str
):
    """
    Background task — creates its own SQLAlchemy session so the request session
    doesn't close before processing finishes.
    """
    db = SessionLocal()
    try:
        # Parse CSV or Excel
        if filename.endswith(".csv"):
            df = pd.read_csv(io.BytesIO(file_content))
        elif filename.endswith((".xlsx", ".xls")):
            df = pd.read_excel(io.BytesIO(file_content))
        else:
            raise ValueError("Unsupported file format — use .csv or .xlsx")

        questions = df.iloc[:, 0].tolist()
        results = batch_process_questionnaire(questions, user_id)

        task = db.query(Task).filter(Task.id == task_id).first()
        if task:
            for r in results:
                db.add(Answer(
                    task_id=task.id,
                    user_id=user_id,
                    question=r["question"],
                    ai_answer=r["answer"],
                    confidence_score=r["confidence"],
                    source_document=r["source_document"],
                    source_snippet=r["source_snippet"]
                ))
            task.status = "completed"
            db.commit()

    except Exception as e:
        task = db.query(Task).filter(Task.id == task_id).first()
        if task:
            task.status = "failed"
            db.commit()
        print(f"[BG] Questionnaire processing failed: {e}")
    finally:
        db.close()


# ─────────────────────────────────────────────
# Task endpoints
# ─────────────────────────────────────────────

@app.get("/api/tasks", tags=["Tasks"])
async def list_tasks(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List all questionnaire tasks for the authenticated user."""
    tasks = db.query(Task).filter(Task.user_id == current_user.id)\
        .order_by(Task.created_at.desc()).all()
    return [
        {
            "id": t.id,
            "task_name": t.task_name,
            "status": t.status,
            "created_at": t.created_at,
            "answer_count": len(t.answers)
        }
        for t in tasks
    ]


@app.get("/api/tasks/{task_id}", response_model=TaskResponse, tags=["Tasks"])
async def get_task(
    task_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get a specific task with all generated answers."""
    task = db.query(Task).filter(
        Task.id == task_id, Task.user_id == current_user.id
    ).first()
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    return TaskResponse.from_orm(task)


@app.put("/api/answers/{answer_id}", tags=["Tasks"])
async def update_answer(
    answer_id: int,
    update_data: AnswerUpdate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Update (manually edit) an AI-generated answer."""
    answer = db.query(Answer).filter(
        Answer.id == answer_id, Answer.user_id == current_user.id
    ).first()
    if not answer:
        raise HTTPException(status_code=404, detail="Answer not found")

    answer.manual_answer = update_data.manual_answer
    answer.has_been_edited = True
    db.commit()
    return AnswerResponse.from_orm(answer)


# ─────────────────────────────────────────────
# Export endpoint
# ─────────────────────────────────────────────

@app.get("/api/tasks/{task_id}/export", tags=["Export"])
async def export_task(
    task_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Export all Q&A pairs for a task as a formatted .docx file.
    Uses a temp file on disk — FileResponse requires a real file path.
    """
    task = db.query(Task).filter(
        Task.id == task_id, Task.user_id == current_user.id
    ).first()
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")

    doc = DocxDocument()

    # Header
    title = doc.add_heading(f"HECVAT Compliance Responses", 0)
    title.alignment = 1  # CENTER
    doc.add_paragraph(f"Task: {task.task_name}")
    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph(f"Powered by AlumniTrust AI — Almabase")
    doc.add_paragraph("")

    for i, answer in enumerate(task.answers, start=1):
        # Question
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"Q{i}. {answer.question}")
        q_run.bold = True
        q_run.font.size = Pt(11)

        # Answer (manual override takes priority)
        final_answer = (
            answer.manual_answer if answer.has_been_edited and answer.manual_answer
            else answer.ai_answer
        )

        a_para = doc.add_paragraph()
        a_para.add_run("Answer: ").bold = True

        if final_answer == "Not found in references":
            run = a_para.add_run(final_answer)
            run.bold = True
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)  # dark red
        else:
            a_para.add_run(final_answer)

        # Metadata line
        meta_para = doc.add_paragraph()
        meta_run = meta_para.add_run(
            f"Confidence: {answer.confidence_score:.1f}%  |  "
            f"Source: {answer.source_document}"
            + (f"  [Edited by reviewer]" if answer.has_been_edited else "")
        )
        meta_run.italic = True
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        doc.add_paragraph("")  # spacer

    # Write to a named temp file (FileResponse needs a real path)
    with tempfile.NamedTemporaryFile(
        suffix=".docx", delete=False, dir=tempfile.gettempdir()
    ) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    safe_name = task.task_name.replace(" ", "_").replace("/", "-")
    return FileResponse(
        path=tmp_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{safe_name}_HECVAT_responses.docx",
        background=BackgroundTasks()  # temp file will be left; OS cleans /tmp
    )


# ─────────────────────────────────────────────
# Health check
# ─────────────────────────────────────────────

@app.get("/api/health", tags=["System"])
async def health_check():
    return {"status": "healthy", "service": "AlumniTrust AI", "version": "1.0.0"}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)
